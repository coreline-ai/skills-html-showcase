#!/usr/bin/env python3
"""
md_to_docx.py — Markdown → DOCX 변환기 (python-docx 전용, 마크다운 라이브러리 미사용)

사용:
  python3 md_to_docx.py <in.md> <out.docx> [--title "제목"] [--subtitle "부제"]

지원 문법:
  - ATX 헤딩 (# ~ ######)
  - 단락(연속 줄은 한 단락으로 병합)
  - 인라인: **굵게**, *기울임*, `인라인코드`, [링크](url)
  - 목록: 순서(1. 2.) / 비순서(- * +), 들여쓰기 중첩
  - GFM 파이프 표 (|---| 구분선 인식, 첫 행 굵게, 'Table Grid' 스타일)
  - 코드블록 (``` ... ```)
  - 인용 (> ...)
  - 수평선 (---  ***  ___)
  - 이미지 ![alt](path)  — 상대경로는 md 파일 기준 해석. 없으면 경고 후 alt 텍스트.

설계 원칙:
  - 외부 마크다운 파서 import 금지. 표준 라이브러리 + python-docx 만.
  - 정규식 기반 라인 파서. 파싱 실패해도 죽지 않고 해당 줄을 일반 단락으로 강등.
  - 한글 폰트(맑은 고딕/Malgun Gothic) 우선 적용, 없는 환경에서도 오류 없이 동작.
"""
import argparse
import os
import re
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[오류] python-docx 가 없습니다.  pip3 install python-docx", file=sys.stderr)
    raise SystemExit(2)

# ----------------------------------------------------------------------------
# 폰트 설정
# ----------------------------------------------------------------------------
KOREAN_FONT = "Pretendard"          # 한글·영문 통일 폰트 (없으면 뷰어가 대체)
KOREAN_FONT_ASCII = "Pretendard"
MONO_FONT = "Consolas"


def _set_korean_font(run, font_name=KOREAN_FONT_ASCII, size=None, bold=None, italic=None, color=None):
    """run 에 한글 폰트(동아시아 글꼴 포함)를 안전하게 적용."""
    try:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        # 동아시아/라틴/복합문자 모두 동일 글꼴로
        rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color
    except Exception:
        # 폰트 적용 실패는 치명적이지 않다 — 기본 글꼴로 진행
        pass


def _configure_styles(doc: Document):
    """Normal/Heading 스타일에 한글 폰트를 박는다."""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = KOREAN_FONT_ASCII
        normal.font.size = Pt(11)
        # R&D·정부 계획서 표준 가독성: 줄간격 1.5, 단락 간격
        try:
            pf = normal.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(6)
        except Exception:
            pass
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
        rfonts.set(qn("w:ascii"), KOREAN_FONT_ASCII)
        rfonts.set(qn("w:hAnsi"), KOREAN_FONT_ASCII)
    except Exception:
        pass
    # 헤딩 스타일에도 한글 글꼴 적용
    for i in range(1, 7):
        try:
            st = doc.styles[f"Heading {i}"]
            st.font.name = KOREAN_FONT_ASCII
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
            rfonts.set(qn("w:ascii"), KOREAN_FONT_ASCII)
            rfonts.set(qn("w:hAnsi"), KOREAN_FONT_ASCII)
        except Exception:
            pass


# ----------------------------------------------------------------------------
# 인라인 파서 — **bold** *italic* `code` [text](url)
# ----------------------------------------------------------------------------
# 토큰 단위로 쪼갠다. 순서가 중요: code 먼저, 그다음 bold, italic, link.
_INLINE_RE = re.compile(
    r"(`[^`]+`)"                       # 인라인 코드
    r"|(\*\*[^*]+\*\*)"                # 굵게 **...**
    r"|(__[^_]+__)"                    # 굵게 __...__
    r"|(\*[^*\n]+\*)"                  # 기울임 *...*
    r"|(_[^_\n]+_)"                    # 기울임 _..._
    r"|(\[[^\]]+\]\([^)]+\))"          # 링크 [text](url)
)


def _add_inline_runs(paragraph, text):
    """인라인 마크업을 파싱해 paragraph 에 run 들을 추가."""
    if text is None:
        return
    pos = 0
    try:
        for m in _INLINE_RE.finditer(text):
            if m.start() > pos:
                _emit_run(paragraph, text[pos:m.start()])
            token = m.group(0)
            if token.startswith("`") and token.endswith("`"):
                run = paragraph.add_run(token[1:-1])
                _set_korean_font(run, font_name=MONO_FONT)
                try:
                    _shade_run(run, "F2F2F2")
                except Exception:
                    pass
            elif (token.startswith("**") and token.endswith("**")) or (
                token.startswith("__") and token.endswith("__")
            ):
                run = paragraph.add_run(token[2:-2])
                _set_korean_font(run, bold=True)
            elif token.startswith("[") and "](" in token:
                # [text](url) — 표시 텍스트만 + 밑줄 파란색
                disp = token[1:token.index("]")]
                run = paragraph.add_run(disp)
                _set_korean_font(run, color=RGBColor(0x1A, 0x5F, 0xB4))
                run.font.underline = True
            else:
                # 기울임 (* 또는 _)
                run = paragraph.add_run(token[1:-1])
                _set_korean_font(run, italic=True)
            pos = m.end()
        if pos < len(text):
            _emit_run(paragraph, text[pos:])
    except Exception:
        # 인라인 파싱 실패 → 원문 그대로
        _emit_run(paragraph, text[pos:])


def _emit_run(paragraph, text):
    if text == "":
        return
    run = paragraph.add_run(text)
    _set_korean_font(run)


def _shade_run(run, hexcolor):
    """run 배경 음영 (인라인 코드용)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    rpr.append(shd)


def _shade_cell(cell, hexcolor):
    """표 셀 배경색."""
    try:
        tcpr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        tcpr.append(shd)
    except Exception:
        pass


# ----------------------------------------------------------------------------
# 블록 파서
# ----------------------------------------------------------------------------
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
HR_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
UL_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
OL_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
IMG_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")
FENCE_RE = re.compile(r"^\s*(`{3,}|~{3,})\s*([\w+-]*)\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")  # |---|:--:|


def _is_table_sep(line):
    """파이프 표의 헤더 구분선인지(--- 와 | 만 포함)."""
    if "|" not in line and "-" not in line:
        return False
    stripped = line.strip()
    if not stripped:
        return False
    # 최소 1개의 '-' 와, '-', ':', '|', 공백 외 문자 없음
    if "-" not in stripped:
        return False
    return all(ch in "-:| \t" for ch in stripped)


def _split_table_row(line):
    """파이프 표 한 행을 셀 리스트로. 양 끝 | 제거, 이스케이프 \\| 보존."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    # 임시 치환으로 이스케이프된 파이프 보호
    s = s.replace(r"\|", "\x00")
    cells = [c.replace("\x00", "|").strip() for c in s.split("|")]
    return cells


class MdParser:
    def __init__(self, doc: Document, base_dir: str):
        self.doc = doc
        self.base_dir = base_dir
        self.warnings = []

    def parse(self, lines):
        i = 0
        n = len(lines)
        para_buf = []  # 단락 누적 버퍼

        def flush_para():
            if para_buf:
                text = " ".join(x.strip() for x in para_buf).strip()
                if text:
                    p = self.doc.add_paragraph()
                    _add_inline_runs(p, text)
                para_buf.clear()

        while i < n:
            raw = lines[i].rstrip("\n")
            line = raw

            # 코드블록 (fenced)
            m = FENCE_RE.match(line)
            if m:
                flush_para()
                fence = m.group(1)[0]
                code_lines = []
                i += 1
                while i < n:
                    cur = lines[i].rstrip("\n")
                    cm = FENCE_RE.match(cur)
                    if cm and cm.group(1)[0] == fence:
                        i += 1
                        break
                    code_lines.append(cur)
                    i += 1
                self._add_code_block(code_lines)
                continue

            # 빈 줄 → 단락 경계
            if line.strip() == "":
                flush_para()
                i += 1
                continue

            # 수평선
            if HR_RE.match(line):
                flush_para()
                self._add_hr()
                i += 1
                continue

            # 헤딩
            hm = HEADING_RE.match(line)
            if hm:
                flush_para()
                level = len(hm.group(1))
                self._add_heading(hm.group(2).strip(), level)
                i += 1
                continue

            # 이미지(단독 줄)
            im = IMG_RE.match(line)
            if im:
                flush_para()
                self._add_image(im.group(1), im.group(2))
                i += 1
                continue

            # 파이프 표: 현재 줄이 |를 포함하고 다음 줄이 구분선
            if "|" in line and i + 1 < n and _is_table_sep(lines[i + 1].rstrip("\n")):
                flush_para()
                i = self._consume_table(lines, i, n)
                continue

            # 인용
            qm = QUOTE_RE.match(line)
            if qm:
                flush_para()
                quote_lines = []
                while i < n:
                    cur = lines[i].rstrip("\n")
                    cqm = QUOTE_RE.match(cur)
                    if cqm:
                        quote_lines.append(cqm.group(1))
                        i += 1
                    elif cur.strip() == "":
                        break
                    else:
                        break
                self._add_quote(quote_lines)
                continue

            # 목록 (순서/비순서)
            if UL_RE.match(line) or OL_RE.match(line):
                flush_para()
                i = self._consume_list(lines, i, n)
                continue

            # 일반 단락 줄 누적
            para_buf.append(line)
            i += 1

        flush_para()

    # --- 블록 빌더들 ---
    def _add_heading(self, text, level):
        try:
            p = self.doc.add_heading(level=min(level, 6))
            p.text = ""  # add_heading 이 빈 run 추가하는 경우 방지
            _add_inline_runs(p, text)
        except Exception:
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            _set_korean_font(run, bold=True, size=16 - level)

    def _add_hr(self):
        p = self.doc.add_paragraph()
        try:
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pbdr.append(bottom)
            ppr.append(pbdr)
        except Exception:
            run = p.add_run("─" * 30)
            _set_korean_font(run)

    def _add_image(self, alt, path):
        # 상대경로는 md 파일 기준 해석
        resolved = path
        if not os.path.isabs(path):
            resolved = os.path.normpath(os.path.join(self.base_dir, path))
        if os.path.exists(resolved):
            try:
                self.doc.add_picture(resolved, width=Inches(5.5))
                if alt:
                    cap = self.doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(alt)
                    _set_korean_font(run, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))
                return
            except Exception as e:
                self.warnings.append(f"이미지 삽입 실패: {resolved} ({e})")
        else:
            self.warnings.append(f"이미지 없음: {resolved} — alt 텍스트로 대체")
        # 없거나 실패 → alt 텍스트
        p = self.doc.add_paragraph()
        run = p.add_run(f"[이미지: {alt or path}]")
        _set_korean_font(run, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    def _add_code_block(self, code_lines):
        p = self.doc.add_paragraph()
        try:
            _shade_paragraph(p, "F2F2F2")
        except Exception:
            pass
        text = "\n".join(code_lines)
        run = p.add_run(text)
        _set_korean_font(run, font_name=MONO_FONT, size=9.5)

    def _add_quote(self, quote_lines):
        text = " ".join(x.strip() for x in quote_lines).strip()
        p = self.doc.add_paragraph()
        try:
            p.style = self.doc.styles["Quote"]
        except Exception:
            pass
        try:
            ppr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "BBBBBB")
            pbdr.append(left)
            ppr.append(pbdr)
        except Exception:
            pass
        _add_inline_runs(p, text)
        for r in p.runs:
            r.font.italic = True

    def _consume_list(self, lines, i, n):
        """연속 목록 블록 처리. 들여쓰기 깊이로 중첩 표현."""
        while i < n:
            raw = lines[i].rstrip("\n")
            if raw.strip() == "":
                # 빈 줄 한 개는 목록 내부 허용하지 않고 종료(단순화)
                break
            um = UL_RE.match(raw)
            om = OL_RE.match(raw)
            if not (um or om):
                break
            if um:
                indent = len(um.group(1).replace("\t", "    "))
                content = um.group(2)
                ordered = False
            else:
                indent = len(om.group(1).replace("\t", "    "))
                content = om.group(3)
                ordered = True
            level = min(indent // 2, 4)
            style = "List Number" if ordered else "List Bullet"
            if level > 0:
                style = f"{style} {level + 1}" if level + 1 <= 3 else style
            try:
                p = self.doc.add_paragraph(style=style)
            except Exception:
                # 해당 스타일이 없으면 기본 + 들여쓰기
                p = self.doc.add_paragraph()
                bullet = ("•", "◦", "▪")[min(level, 2)] if not ordered else "•"
                run = p.add_run("    " * level + (f"{bullet} "))
                _set_korean_font(run)
            _add_inline_runs(p, content)
            i += 1
        return i

    def _consume_table(self, lines, i, n):
        """파이프 표 처리. i 는 헤더 행, i+1 은 구분선."""
        header = _split_table_row(lines[i].rstrip("\n"))
        ncols = len(header)
        body = []
        j = i + 2  # 구분선 건너뜀
        while j < n:
            cur = lines[j].rstrip("\n")
            if "|" not in cur or cur.strip() == "":
                break
            if _is_table_sep(cur):
                j += 1
                continue
            cells = _split_table_row(cur)
            # 컬럼 수 보정
            if len(cells) < ncols:
                cells += [""] * (ncols - len(cells))
            elif len(cells) > ncols:
                cells = cells[:ncols]
            body.append(cells)
            j += 1

        try:
            table = self.doc.add_table(rows=1, cols=ncols)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            hdr_cells = table.rows[0].cells
            for c, htext in enumerate(header):
                hdr_cells[c].text = ""
                para = hdr_cells[c].paragraphs[0]
                _add_inline_runs(para, htext)
                for r in para.runs:
                    r.font.bold = True
                _shade_cell(hdr_cells[c], "D9E2F3")
            for row in body:
                cells = table.add_row().cells
                for c, ctext in enumerate(row):
                    cells[c].text = ""
                    para = cells[c].paragraphs[0]
                    _add_inline_runs(para, ctext)
            self.doc.add_paragraph()  # 표 뒤 간격
        except Exception as e:
            self.warnings.append(f"표 변환 실패 → 텍스트로 강등 ({e})")
            for row in [header] + body:
                p = self.doc.add_paragraph()
                _add_inline_runs(p, " | ".join(row))
        return j


def _shade_paragraph(p, hexcolor):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    ppr.append(shd)


def main() -> int:
    ap = argparse.ArgumentParser(description="Markdown → DOCX 변환 (python-docx 전용)")
    ap.add_argument("input", nargs="?", help="입력 .md 파일")
    ap.add_argument("output", nargs="?", help="출력 .docx 파일")
    ap.add_argument("--title", default="", help="문서 표지 제목")
    ap.add_argument("--subtitle", default="", help="문서 표지 부제")
    args = ap.parse_args()

    if not args.input or not args.output:
        ap.print_usage()
        print("\n예) python3 md_to_docx.py research-report.md research-report.docx --title \"리서치 보고서\"",
              file=sys.stderr)
        return 1

    if not os.path.exists(args.input):
        print(f"[오류] 입력 파일 없음: {args.input}", file=sys.stderr)
        return 1

    try:
        with open(args.input, encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        print(f"[오류] 파일 읽기 실패: {e}", file=sys.stderr)
        return 1

    doc = Document()
    _configure_styles(doc)

    # 표지 제목/부제
    if args.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(args.title)
        _set_korean_font(run, size=22, bold=True)
    if args.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(args.subtitle)
        _set_korean_font(run, size=13, color=RGBColor(0x66, 0x66, 0x66))
    if args.title or args.subtitle:
        doc.add_paragraph()

    parser = MdParser(doc, base_dir=os.path.dirname(os.path.abspath(args.input)))
    parser.parse(content.split("\n"))

    out_dir = os.path.dirname(os.path.abspath(args.output))
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    try:
        doc.save(args.output)
    except Exception as e:
        print(f"[오류] DOCX 저장 실패: {e}", file=sys.stderr)
        return 1

    print(f"[완료] DOCX 생성: {args.output}")
    if parser.warnings:
        print(f"[경고 {len(parser.warnings)}건]")
        for w in parser.warnings:
            print(f"  ⚠ {w}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
